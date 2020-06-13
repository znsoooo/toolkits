# build 20200305

import os
import re
import urllib
import urllib.request
import html

import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def ReadPassage(url):
    request  = urllib.request.Request(url)
    response = urllib.request.urlopen(request, timeout=5)
    content  = response.read().decode()
    content2 = html.unescape(content.replace('\n',''))

    date    = re.findall(r'var t=".*?",n=".*?",s="(.*?)"', content2)[0]
    date    = date.replace('-', '') 
    profile = re.findall(r'<strong class="profile_nickname">(.*?)</strong>', content2)[0]
    author  = re.findall(r'<meta name="author" content="(.*?)" />', content2)[0]
    title   = re.findall(r'<meta property="og:title" content="(.*?)" />', content2)[0]

    passage = re.findall(r'<div class="rich_media_content ".*?>(.*?)</div>', content2)[0]

    print(date, profile, author, title)

    # 去除文件夹命名的非法字符
    folder = '%s %s %s %s'%(date, profile, author, title)
    for s in '\t\\/:*?"<>|':
        folder = folder.replace(s, '')
    os.makedirs(folder, exist_ok=1)

    return folder, profile, author, title, passage


def DownloadImage(img_url, path):
    img_url  = img_url.rsplit('/',1)[0]+'/'
    request  = urllib.request.Request(img_url)
    response = urllib.request.urlopen(request, timeout=5)
    data     = response.read()
    with open(path, 'wb') as f:
        f.write(data)


def SaveDocument(save_file, folder, author, title, paragraphs):
    doc = docx.Document()
    doc.add_heading(title, 0)
    p = doc.add_paragraph(author)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cnt = 0
    for paragraph in paragraphs:
        doc.add_paragraph(re.sub(r'<.*?>', '', paragraph)) # TODO 保留加粗等格式
        imgs = re.findall(r'<img.*?data-src="(.*?)".*?/>', paragraph)
        for img in imgs:
            doc.add_picture('%s/%s.jpg'%(folder, cnt), width=Inches(6))
            cnt += 1
    doc.save(save_file)


def SavePassage(url):
    folder, profile, author, title, passage = ReadPassage(url)
    passage = passage.replace('<p ', '<section ').replace('</p>', '</section>')
    paragraphs = re.findall(r'<section.*?>(.*?)</section>', passage)

    # 获取正文中的文本和提取图片链接
    img_urls = []
    passage2 = []
    for paragraph in paragraphs:
        passage2.append(re.sub(r'<.*?>', '', paragraph))
        img_urls += re.findall(r'<img.*?data-src="(.*?)".*?/>', paragraph)

    # 下载文章中的所有图片
    for i, img_url in enumerate(img_urls):
        DownloadImage(img_url, '%s/%s.jpg'%(folder, i))

    # 保存正文为纯文本文件
    with open('%s/%s.txt'%(folder, folder), 'w', encoding='u8') as f:
        f.write('\n'.join(passage2))

    # 保存文章为Word文档
    SaveDocument('%s/%s.docx'%(folder, folder), folder, profile+' '+author, title, paragraphs)


##url = 'https://mp.weixin.qq.com/s/k87AxlcmI9pomFjgblpA5g'
##
##SavePassage(url)


##SavePassage('https://mp.weixin.qq.com/s/OK4jaiAPMF58ycI25KU4zQ')
##SavePassage('https://mp.weixin.qq.com/s/t9FZUXziG4AqcmEpLibrUg')
##SavePassage('https://mp.weixin.qq.com/s/SnhVPSh4i5dL0BJD5kRyeQ')
##SavePassage('https://mp.weixin.qq.com/s/vUdmTAUOBan1CDntcfOWvA')
##SavePassage('https://mp.weixin.qq.com/s/IMpUyvjXzK6hXtU6jqst_A')
##SavePassage('https://mp.weixin.qq.com/s/SqsjArgDKD-GZ71L4Qsa_Q')
##SavePassage('https://mp.weixin.qq.com/s/1D_GkQs2oQ_1vYS2dqZ0PA')
##SavePassage('https://mp.weixin.qq.com/s/MMIXRL2KxMg3xKVczFYDFA')
##SavePassage('https://mp.weixin.qq.com/s/pkkFIhYiKMg2vugsgurEew')
##SavePassage('https://mp.weixin.qq.com/s/wW9JuTvA74GymOCFF2u_4w')
##SavePassage('https://mp.weixin.qq.com/s/nf51GRLFmMIUYohSpUilIg')
##SavePassage('https://mp.weixin.qq.com/s/k87AxlcmI9pomFjgblpA5g')

##SavePassage('https://mp.weixin.qq.com/s?__biz=MzIzMTc2OTYxMg==&tempkey=MTA2MF96RGRLV0lZbldVWXJCNmErdGxzWkZzYTZkaTBQdE01OFcyZ3NoVDFPSUtLVnh6cFkycEk4aG5GR2JQT1FOVW1qZmhOU2lWd2dyd0JBSF9LTkU5SFVuTndpUFZpWHNqZF94a0dGRlotVUVXbnFQTF8tYWRsaGJpd0lFcTd2QTl5VDBfSkctaWRUdnZyVUJseGZmVnZ5cDF6b0c4QzZsOG9tb1JXY1dBfn4%3D&chksm=689e50b25fe9d9a4da645388423ed6efaadb9d6c08319ceae63cb34e7451bc31cd2fadce0534#rd')

##s = '''var msgList = '{&quot;list&quot;:[{&quot;comm_msg_info&quot;:{&quot;id&quot;:1000000027,&quot;type&quot;:49,&quot;datetime&quot;:1588949416,&quot;fakeid&quot;:&quot;3231769612&quot;,&quot;status&quot;:2,&quot;content&quot;:&quot;&quot;},&quot;app_msg_ext_info&quot;:{&quot;title&quot;:&quot;战疫·增强免疫力&nbsp;|&nbsp;系统产品事业部成功举办四月份全民运动健身活动&quot;,&quot;digest&quot;:&quot;全民健身，国家战略；疫情之下，尤为重要。在十二所工会的指导下，系统产品事业部“战疫•增强免疫力”全民运动健身&quot;,&quot;content&quot;:&quot;&quot;,&quot;fileid&quot;:100000722,&quot;content_url&quot;:&quot;http://mp.weixin.qq.com/s?__biz=MzIzMTc2OTYxMg==&amp;amp;mid=2247484395&amp;amp;idx=1&amp;amp;sn=9558644375a759c287e9847ff2d9f80f&amp;amp;chksm=e89e5088dfe9d99e681d3cc3999ae9a03fafa0aa8dbebc80d0fd7f0b24ae079fd83856a01f47&amp;amp;scene=27#wechat_redirect&quot;,&quot;source_url&quot;:&quot;&quot;,&quot;cover&quot;:&quot;http://mmbiz.qpic.cn/mmbiz_jpg/Z2f4yGjIu86LEJEHejNTicww1PCRz7YM7CqdQGXr8setof96MplTEAFUiajcDmxIOQJ8ic7y3pm1ibWPyCbpV32OrQ/0?wx_fmt=jpeg&quot;,&quot;subtype&quot;:9,&quot;is_multi&quot;:1,&quot;multi_app_msg_item_list&quot;:[{&quot;title&quot;:&quot;春风十里，不如健身的你&quot;,&quot;digest&quot;:&quot;每天锻炼一小时，健康工作五十年，幸福生活一辈子。&quot;,&quot;content&quot;:&quot;&quot;,&quot;fileid&quot;:0,&quot;content_url&quot;:&quot;http://mp.weixin.qq.com/s?__biz=MzIzMTc2OTYxMg==&amp;amp;mid=2247484395&amp;amp;idx=2&amp;amp;sn=57fe8090c6b239627a241f1c541677e1&amp;amp;chksm=e89e5088dfe9d99e609ac0bfe8629018e51409c541978f4845f8d74a18386420412e1266d2fc&amp;amp;scene=27#wechat_redirect&quot;,&quot;source_url&quot;:&quot;&quot;,&quot;cover&quot;:&quot;http://mmbiz.qpic.cn/mmbiz_jpg/Z2f4yGjIu86LEJEHejNTicww1PCRz7YM7uibcGLLYLoJavhwKFQwNa3VG6dGqib23QfXQGFpegJth1CIkVszkU7RQ/0?wx_fmt=jpeg&quot;,&quot;author&quot;:&quot;&quot;,&quot;copyright_stat&quot;:100,&quot;del_flag&quot;:1,&quot;item_show_type&quot;:5,&quot;audio_fileid&quot;:0,&quot;duration&quot;:0,&quot;play_url&quot;:&quot;&quot;,&quot;malicious_title_reason_id&quot;:0,&quot;malicious_content_type&quot;:0},{&quot;title&quot;:&quot;系统产品事业部五月生日祝福！&quot;,&quot;digest&quot;:&quot;点击阅读原文接收生日祝福贺卡吧！&quot;,&quot;content&quot;:&quot;&quot;,&quot;fileid&quot;:0,&quot;content_url&quot;:&quot;http://mp.weixin.qq.com/s?__biz=MzIzMTc2OTYxMg==&amp;amp;mid=2247484395&amp;amp;idx=3&amp;amp;sn=483b5aa2e26303a913425fb64a823fa6&amp;amp;chksm=e89e5088dfe9d99efab4112cafba31d1e752755f69fb8b5fef965e5e30f5f9b1075f93607538&amp;amp;scene=27#wechat_redirect&quot;,&quot;source_url&quot;:&quot;http://u7160780.viewer.maka.im/k/U3YC1Z7RW7160780?t=1588941159785&quot;,&quot;cover&quot;:&quot;http://mmbiz.qpic.cn/mmbiz_jpg/Z2f4yGjIu87v7tjUaxsg79kcjnVcFTl7xickDpoU1jXw8c8dg75y8GUBABmpibc4puDO8wAKGOemkNuQoCDalcMg/0?wx_fmt=jpeg&quot;,&quot;author&quot;:&quot;系统产品事业部&quot;,&quot;copyright_stat&quot;:100,&quot;del_flag&quot;:1,&quot;item_show_type&quot;:0,&quot;audio_fileid&quot;:0,&quot;duration&quot;:0,&quot;play_url&quot;:&quot;&quot;,&quot;malicious_title_reason_id&quot;:0,&quot;malicious_content_type&quot;:0}],&quot;author&quot;:&quot;李世先&quot;,&quot;copyright_stat&quot;:11,&quot;duration&quot;:0,&quot;del_flag&quot;:1,&quot;item_show_type&quot;:0,&quot;audio_fileid&quot;:0,&quot;play_url&quot;:&quot;&quot;,&quot;malicious_title_reason_id&quot;:0,&quot;malicious_content_type&quot;:0}},{&quot;comm_msg_info&quot;:{&quot;id&quot;:1000000026,&quot;type&quot;:49,&quot;datetime&quot;:1588567043,&quot;fakeid&quot;:&quot;3231769612&quot;,&quot;status&quot;:2,&quot;content&quot;:&quot;&quot;},&quot;app_msg_ext_info&quot;:{&quot;title&quot;:&quot;话青年使命&nbsp;•&nbsp;展青春担当——四室与六室团支部“五四”主题团日共建活动&quot;,&quot;digest&quot;:&quot;5月4日，四室和六室团支部，在线上举行纪念五四运动101周年暨“话青年使命，展青春担当”主题团日活动。&quot;,&quot;content&quot;:&quot;&quot;,&quot;fileid&quot;:100000700,&quot;content_url&quot;:&quot;http://mp.weixin.qq.com/s?__biz=MzIzMTc2OTYxMg==&amp;amp;mid=2247484366&amp;amp;idx=1&amp;amp;sn=ece4afde1277d385d43d03249f947814&amp;amp;chksm=e89e50addfe9d9bb6cca438037e1ef6f3138e599510e732f8ed691925b045495e9fc0a6e97db&amp;amp;scene=27#wechat_redirect&quot;,&quot;source_url&quot;:&quot;&quot;,&quot;cover&quot;:&quot;http://mmbiz.qpic.cn/mmbiz_jpg/Z2f4yGjIu86vMDKsXI4mJIUVVAf2iavww04Ysa5Cic50vevVUDbA035T3LobUl1uATmnS8kdeJXyQiaWbjjJbvyOw/0?wx_fmt=jpeg&quot;,&quot;subtype&quot;:9,&quot;is_multi&quot;:0,&quot;multi_app_msg_item_list&quot;:[],&quot;author&quot;:&quot;杨诚&quot;,&quot;copyright_stat&quot;:11,&quot;duration&quot;:0,&quot;del_flag&quot;:1,&quot;item_show_type&quot;:0,&quot;audio_fileid&quot;:0,&quot;play_url&quot;:&quot;&quot;,&quot;malicious_title_reason_id&quot;:0,&quot;malicious_content_type&quot;:0}},{&quot;comm_msg_info&quot;:{&quot;id&quot;:1000000025,&quot;type&quot;:49,&quot;datetime&quot;:1585672116,&quot;fakeid&quot;:&quot;3231769612&quot;,&quot;status&quot;:2,&quot;content&quot;:&quot;&quot;},&quot;app_msg_ext_info&quot;:{&quot;title&quot;:&quot;战疫&nbsp;·&nbsp;增强免疫力&nbsp;|&nbsp;系统产品事业部开展全民运动健身活动&quot;,&quot;digest&quot;:&quot;全民健身，国家战略；疫情之下，尤为重要。在十二所工会的指导下，系统产品事业部“战疫•增强免疫力”全民运动健身&quot;,&quot;content&quot;:&quot;&quot;,&quot;fileid&quot;:100000688,&quot;content_url&quot;:&quot;http://mp.weixin.qq.com/s?__biz=MzIzMTc2OTYxMg==&amp;amp;mid=2247484346&amp;amp;idx=1&amp;amp;sn=746d917c06688d21d99744dbfbb39dfa&amp;amp;chksm=e89e50d9dfe9d9cf09861df79b18bd766a1fca5d01b29a774bb660b659932990230e217ae4cf&amp;amp;scene=27#wechat_redirect&quot;,&quot;source_url&quot;:&quot;&quot;,&quot;cover&quot;:&quot;http://mmbiz.qpic.cn/mmbiz_jpg/Z2f4yGjIu87Fc4JvzVkicHNQUhlT9BvguOnYV5McC9OHvjCg7e8mgAluaNwiboU3XcpKBUk5ZAYqlwoYAjSnk10A/0?wx_fmt=jpeg&quot;,&quot;subtype&quot;:9,&quot;is_multi&quot;:1,&quot;multi_app_msg_item_list&quot;:[{&quot;title&quot;:&quot;系统产品事业部四月生日祝福！&quot;,&quot;digest&quot;:&quot;点击阅读原文接收生日祝福贺卡吧！&quot;,&quot;content&quot;:&quot;&quot;,&quot;fileid&quot;:0,&quot;content_url&quot;:&quot;http://mp.weixin.qq.com/s?__biz=MzIzMTc2OTYxMg==&amp;amp;mid=2247484346&amp;amp;idx=2&amp;amp;sn=27a691c138b5d1deaf50cf0c1f6ad0fc&amp;amp;chksm=e89e50d9dfe9d9cf42068d5cb1ce0e8b99d3af39899cc95083d64a1d973c4e37335566e87e34&amp;amp;scene=27#wechat_redirect&quot;,&quot;source_url&quot;:&quot;http://u7160780.viewer.maka.im/k/ELFKOKRJW7160780?t=1585655323499&quot;,&quot;cover&quot;:&quot;http://mmbiz.qpic.cn/mmbiz_jpg/Z2f4yGjIu87v7tjUaxsg79kcjnVcFTl7xickDpoU1jXw8c8dg75y8GUBABmpibc4puDO8wAKGOemkNuQoCDalcMg/0?wx_fmt=jpeg&quot;,&quot;author&quot;:&quot;系统产品事业部&quot;,&quot;copyright_stat&quot;:100,&quot;del_flag&quot;:1,&quot;item_show_type&quot;:0,&quot;audio_fileid&quot;:0,&quot;duration&quot;:0,&quot;play_url&quot;:&quot;&quot;,&quot;malicious_title_reason_id&quot;:0,&quot;malicious_content_type&quot;:0}],&quot;author&quot;:&quot;系统产品事业部&quot;,&quot;copyright_stat&quot;:11,&quot;duration&quot;:0,&quot;del_flag&quot;:1,&quot;item_show_type&quot;:0,&quot;audio_fileid&quot;:0,&quot;play_url&quot;:&quot;&quot;,&quot;malicious_title_reason_id&quot;:0,&quot;malicious_content_type&quot;:0}},{&quot;comm_msg_info&quot;:{&quot;id&quot;:1000000024,&quot;type&quot;:49,&quot;datetime&quot;:1583798420,&quot;fakeid&quot;:&quot;3231769612&quot;,&quot;status&quot;:2,&quot;content&quot;:&quot;&quot;},&quot;app_msg_ext_info&quot;:{&quot;title&quot;:&quot;“防疫不停工，你我共学习”第四期：提升企业创新力的学习交流&quot;,&quot;digest&quot;:&quot;“学而不思则罔，思而不学则殆。&quot;,&quot;content&quot;:&quot;&quot;,&quot;fileid&quot;:100000659,&quot;content_url&quot;:&quot;http://mp.weixin.qq.com/s?__biz=MzIzMTc2OTYxMg==&amp;amp;mid=2247484335&amp;amp;idx=1&amp;amp;sn=3781eb93acf05c0c50ed893a18ee3315&amp;amp;chksm=e89e50ccdfe9d9da50786578dd512fe266550e14f7e02230484fcb89767dda49b515d035e8b1&amp;amp;scene=27#wechat_redirect&quot;,&quot;source_url&quot;:&quot;&quot;,&quot;cover&quot;:&quot;http://mmbiz.qpic.cn/mmbiz_jpg/Z2f4yGjIu84iafWCzr4U5raKlAsWGNAalqvqW2hFv1kYCNTibyIibfFpJbOew1D1dw9Kg0RjSTksibzQBvg1yfXrhg/0?wx_fmt=jpeg&quot;,&quot;subtype&quot;:9,&quot;is_multi&quot;:1,&quot;multi_app_msg_item_list&quot;:[{&quot;title&quot;:&quot;“防疫不停工，你我共学习”第五期：人工智能概述&quot;,&quot;digest&quot;:&quot;登鹳雀楼白日依山尽，黄河入海流。欲穷千里目，更上一层楼。人工智能概述按照2020年六室防疫期间培训策划及要求&quot;,&quot;content&quot;:&quot;&quot;,&quot;fileid&quot;:0,&quot;content_url&quot;:&quot;http://mp.weixin.qq.com/s?__biz=MzIzMTc2OTYxMg==&amp;amp;mid=2247484335&amp;amp;idx=2&amp;amp;sn=ec95bc3eff693aeea38dbe216bf9c2dc&amp;amp;chksm=e89e50ccdfe9d9da667e48ce7a87403f6150c9fc1e56aae80e7adbe082fdf057baabdc7dbe7c&amp;amp;scene=27#wechat_redirect&quot;,&quot;source_url&quot;:&quot;&quot;,&quot;cover&quot;:&quot;http://mmbiz.qpic.cn/mmbiz_jpg/Z2f4yGjIu85bHHOBEKoOlpibia47kbFCiaraYRjric1EuIfxL1K2roNvhcXaNvawyDa5lFdtvibkbB490Ufw8OxzQDg/0?wx_fmt=jpeg&quot;,&quot;author&quot;:&quot;李世先&quot;,&quot;copyright_stat&quot;:11,&quot;del_flag&quot;:1,&quot;item_show_type&quot;:0,&quot;audio_fileid&quot;:0,&quot;duration&quot;:0,&quot;play_url&quot;:&quot;&quot;,&quot;malicious_title_reason_id&quot;:0,&quot;malicious_content_type&quot;:0},{&quot;title&quot;:&quot;“防疫不停工，你我共学习”第六期：高性能PSOC专题技术交流&quot;,&quot;digest&quot;:&quot;冬夜读书示子聿古人学问无遗力，少壮工夫老始成。纸上得来终觉浅，绝知此事要躬行。高性能PSOC专题技术交流庚子&quot;,&quot;content&quot;:&quot;&quot;,&quot;fileid&quot;:100000643,&quot;content_url&quot;:&quot;http://mp.weixin.qq.com/s?__biz=MzIzMTc2OTYxMg==&amp;amp;mid=2247484335&amp;amp;idx=3&amp;amp;sn=3a35913da90b75621a27c9481b48480b&amp;amp;chksm=e89e50ccdfe9d9da37ad93acf8a9b895b4f6632c9ff6567937883b81f78445e5b81aab9f52e3&amp;amp;scene=27#wechat_redirect&quot;,&quot;source_url&quot;:&quot;&quot;,&quot;cover&quot;:&quot;http://mmbiz.qpic.cn/mmbiz_jpg/Z2f4yGjIu87NofkpCMQbN0TiaWuzdOAhqe57YBKzAG3IXh4YNzTl2X1T3s9WYoeWI6PBSKeHF9VJXWuoE3t2Sww/0?wx_fmt=jpeg&quot;,&quot;author&quot;:&quot;张红&quot;,&quot;copyright_stat&quot;:11,&quot;del_flag&quot;:1,&quot;item_show_type&quot;:0,&quot;audio_fileid&quot;:0,&quot;duration&quot;:0,&quot;play_url&quot;:&quot;&quot;,&quot;malicious_title_reason_id&quot;:0,&quot;malicious_content_type&quot;:0},{&quot;title&quot;:&quot;“防疫不停工，你我共学习”第七期：惯性器件专题交流&quot;,&quot;digest&quot;:&quot;这次培训，有幸请来了事业部技术大牛柳柱，小伙伴们也早早进入了会议系统等待听讲。柳主任讲课风趣幽默、深入浅出，为了让大家能更好的了解控制系统工作原理。&quot;,&quot;content&quot;:&quot;&quot;,&quot;fileid&quot;:100000682,&quot;content_url&quot;:&quot;http://mp.weixin.qq.com/s?__biz=MzIzMTc2OTYxMg==&amp;amp;mid=2247484335&amp;amp;idx=4&amp;amp;sn=31a1272fce7803c6510b20c6241eccf7&amp;amp;chksm=e89e50ccdfe9d9da7cc37adf4642a88fe26288f7e0a513840a15d1dab9f95f8acd2b17b4b7df&amp;amp;scene=27#wechat_redirect&quot;,&quot;source_url&quot;:&quot;&quot;,&quot;cover&quot;:&quot;http://mmbiz.qpic.cn/mmbiz_jpg/Z2f4yGjIu85LVlPibGUfibNa7bSuLWKBavZc9LsFfrtukrNKPW2dSupmCDgubLG40TYQjqAWUPWvRbBic57E5JX9g/0?wx_fmt=jpeg&quot;,&quot;author&quot;:&quot;二部一组&quot;,&quot;copyright_stat&quot;:11,&quot;del_flag&quot;:1,&quot;item_show_type&quot;:0,&quot;audio_fileid&quot;:0,&quot;duration&quot;:0,&quot;play_url&quot;:&quot;&quot;,&quot;malicious_title_reason_id&quot;:0,&quot;malicious_content_type&quot;:0}],&quot;author&quot;:&quot;王大林&quot;,&quot;copyright_stat&quot;:11,&quot;duration&quot;:0,&quot;del_flag&quot;:1,&quot;item_show_type&quot;:0,&quot;audio_fileid&quot;:0,&quot;play_url&quot;:&quot;&quot;,&quot;malicious_title_reason_id&quot;:0,&quot;malicious_content_type&quot;:0}},{&quot;comm_msg_info&quot;:{&quot;id&quot;:1000000023,&quot;type&quot;:49,&quot;datetime&quot;:1583624075,&quot;fakeid&quot;:&quot;3231769612&quot;,&quot;status&quot;:2,&quot;content&quot;:&quot;&quot;},&quot;app_msg_ext_info&quot;:{&quot;title&quot;:&quot;系统产品事业部祝女神们节日快乐！&quot;,&quot;digest&quot;:&quot;在建设航天强国的征途中，您是坚韧的半边天，靓丽的风景线，愿您不忘初心携梦前行，为探索浩瀚宇宙，支撑世界一流军队建设，绽放出更加璀璨夺目的光辉！&quot;,&quot;content&quot;:&quot;&quot;,&quot;fileid&quot;:0,&quot;content_url&quot;:&quot;http://mp.weixin.qq.com/s?__biz=MzIzMTc2OTYxMg==&amp;amp;mid=2247484321&amp;amp;idx=1&amp;amp;sn=f1318b67918c8863fbbf6dc7f20d773a&amp;amp;chksm=e89e50c2dfe9d9d4cedfe4bb31a5893f15bf89b47cb0472c4560b0d11e8a7be585ceedd7a8b4&amp;amp;scene=27#wechat_redirect&quot;,&quot;source_url&quot;:&quot;&quot;,&quot;cover&quot;:&quot;http://mmbiz.qpic.cn/mmbiz_jpg/Z2f4yGjIu87v7tjUaxsg79kcjnVcFTl7gPMs3SOX5GDSKonducib70IASAIMdXmYhxvcwc1SFm2GE5t1VROdEpA/0?wx_fmt=jpeg&quot;,&quot;subtype&quot;:9,&quot;is_multi&quot;:1,&quot;multi_app_msg_item_list&quot;:[{&quot;title&quot;:&quot;系统产品事业部三月生日祝福！&quot;,&quot;digest&quot;:&quot;点击阅读原文接收生日祝福贺卡吧！&quot;,&quot;content&quot;:&quot;&quot;,&quot;fileid&quot;:0,&quot;content_url&quot;:&quot;http://mp.weixin.qq.com/s?__biz=MzIzMTc2OTYxMg==&amp;amp;mid=2247484321&amp;amp;idx=2&amp;amp;sn=0cf142d841affd0057a03d60b01b33fe&amp;amp;chksm=e89e50c2dfe9d9d4487210b1a41b3f8627a9ce5e25479220f340dce8e7fadf7159f9e012753b&amp;amp;scene=27#wechat_redirect&quot;,&quot;source_url&quot;:&quot;http://u7160780.viewer.maka.im/k/V9HEYO1SW7160780?t=1583582957254&quot;,&quot;cover&quot;:&quot;http://mmbiz.qpic.cn/mmbiz_jpg/Z2f4yGjIu87v7tjUaxsg79kcjnVcFTl7xickDpoU1jXw8c8dg75y8GUBABmpibc4puDO8wAKGOemkNuQoCDalcMg/0?wx_fmt=jpeg&quot;,&quot;author&quot;:&quot;系统产品事业部&quot;,&quot;copyright_stat&quot;:100,&quot;del_flag&quot;:1,&quot;item_show_type&quot;:0,&quot;audio_fileid&quot;:0,&quot;duration&quot;:0,&quot;play_url&quot;:&quot;&quot;,&quot;malicious_title_reason_id&quot;:0,&quot;malicious_content_type&quot;:0}],&quot;author&quot;:&quot;系统产品事业部&quot;,&quot;copyright_stat&quot;:11,&quot;duration&quot;:0,&quot;del_flag&quot;:1,&quot;item_show_type&quot;:0,&quot;audio_fileid&quot;:0,&quot;play_url&quot;:&quot;&quot;,&quot;malicious_title_reason_id&quot;:0,&quot;malicious_content_type&quot;:0}},{&quot;comm_msg_info&quot;:{&quot;id&quot;:1000000022,&quot;type&quot;:49,&quot;datetime&quot;:1583021713,&quot;fakeid&quot;:&quot;3231769612&quot;,&quot;status&quot;:2,&quot;content&quot;:&quot;&quot;},&quot;app_msg_ext_info&quot;:{&quot;title&quot;:&quot;&amp;quot;疫&amp;quot;往无前&nbsp;|&nbsp;抗战新冠临危不乱，科学复工有条不紊&quot;,&quot;digest&quot;:&quot;“&nbsp;系统产品事业部全体同仁：自新冠肺炎爆发以来，在所党委的坚强领导下，在所疫情防控指挥部直接部署指导下，事业&quot;,&quot;content&quot;:&quot;&quot;,&quot;fileid&quot;:100000625,&quot;content_url&quot;:&quot;http://mp.weixin.qq.com/s?__biz=MzIzMTc2OTYxMg==&amp;amp;mid=2247484282&amp;amp;idx=1&amp;amp;sn=49bdc07750c50297f925b96260677aab&amp;amp;chksm=e89e5019dfe9d90fcc54afea117cd12432ba104bf93900533d5910be9427aa5039c9ee37c6da&amp;amp;scene=27#wechat_redirect&quot;,&quot;source_url&quot;:&quot;&quot;,&quot;cover&quot;:&quot;http://mmbiz.qpic.cn/mmbiz_jpg/Z2f4yGjIu86oGv9lFn90JIEto9dzuPu0nCqHicQsYXTQwGeFSLT2ogvlc22AXT6pyx8uZKCiczhMGMWUaUic9A2Bg/0?wx_fmt=jpeg&quot;,&quot;subtype&quot;:9,&quot;is_multi&quot;:1,&quot;multi_app_msg_item_list&quot;:[{&quot;title&quot;:&quot;“疫”呼百应&nbsp;|&nbsp;一名湖北籍职工的所想所感&quot;,&quot;digest&quot;:&quot;庚子鼠年的春节可能是这么多年来最惊心动魄的一个，一切发生得出乎意料之外，打得大家措手不及。疫情一天天的快速增&quot;,&quot;content&quot;:&quot;&quot;,&quot;fileid&quot;:0,&quot;content_url&quot;:&quot;http://mp.weixin.qq.com/s?__biz=MzIzMTc2OTYxMg==&amp;amp;mid=2247484282&amp;amp;idx=2&amp;amp;sn=a504b991d7f36136aa1984decfd2de35&amp;amp;chksm=e89e5019dfe9d90f8ad0378470b7d0af657c1aa1dc983f2b69f775db9494c2871f6021ad9f23&amp;amp;scene=27#wechat_redirect&quot;,&quot;source_url&quot;:&quot;&quot;,&quot;cover&quot;:&quot;http://mmbiz.qpic.cn/mmbiz_jpg/Z2f4yGjIu86oGv9lFn90JIEto9dzuPu0dicRngOHt8s75HicVPrFEzhk1BpAAFiaO7c1m9f8xGvJpMR8vsvMONknA/0?wx_fmt=jpeg&quot;,&quot;author&quot;:&quot;&quot;,&quot;copyright_stat&quot;:100,&quot;del_flag&quot;:1,&quot;item_show_type&quot;:0,&quot;audio_fileid&quot;:0,&quot;duration&quot;:0,&quot;play_url&quot;:&quot;&quot;,&quot;malicious_title_reason_id&quot;:0,&quot;malicious_content_type&quot;:0},{&quot;title&quot;:&quot;“疫”言难尽&nbsp;|&nbsp;记一个未返京共产党员的心声&quot;,&quot;digest&quot;:&quot;时光匆匆，这个冬天不知道从什么时候开始下雪，也记不得刮过了几场肆虐的狂风，2019年的冬天过得是沉重的，新型&quot;,&quot;content&quot;:&quot;&quot;,&quot;fileid&quot;:0,&quot;content_url&quot;:&quot;http://mp.weixin.qq.com/s?__biz=MzIzMTc2OTYxMg==&amp;amp;mid=2247484282&amp;amp;idx=3&amp;amp;sn=d34d8b30fbc5f71fe45836409b2335c6&amp;amp;chksm=e89e5019dfe9d90fd6efcea07bb55dc6758130ecbaa3f206cf3246ed86959ec6e7b7b01ae907&amp;amp;scene=27#wechat_redirect&quot;,&quot;source_url&quot;:&quot;&quot;,&quot;cover&quot;:&quot;http://mmbiz.qpic.cn/mmbiz_jpg/Z2f4yGjIu86oGv9lFn90JIEto9dzuPu0xffqUMABgoic5XHKpdKh96GMFZggOG62hYiaLiaSu3Uef2cpfq7K0nAeA/0?wx_fmt=jpeg&quot;,&quot;author&quot;:&quot;&quot;,&quot;copyright_stat&quot;:100,&quot;del_flag&quot;:1,&quot;item_show_type&quot;:0,&quot;audio_fileid&quot;:0,&quot;duration&quot;:0,&quot;play_url&quot;:&quot;&quot;,&quot;malicious_title_reason_id&quot;:0,&quot;malicious_content_type&quot;:0},{&quot;title&quot;:&quot;“疫”苦思甜&nbsp;|&nbsp;岂曰无衣，与子同袍&quot;,&quot;digest&quot;:&quot;春节，是中国的传统节日，不论是在外工作的人，还是异国他乡的游子们，无论有多远，都极其渴望回家过年，与家人团圆&quot;,&quot;content&quot;:&quot;&quot;,&quot;fileid&quot;:100000629,&quot;content_url&quot;:&quot;http://mp.weixin.qq.com/s?__biz=MzIzMTc2OTYxMg==&amp;amp;mid=2247484282&amp;amp;idx=4&amp;amp;sn=3274bc8d752af3e9e2e28740f16bae75&amp;amp;chksm=e89e5019dfe9d90f8ab9da2a43f350fe44c5483b5601d1028f4050a683f7556bdf69101bbd5a&amp;amp;scene=27#wechat_redirect&quot;,&quot;source_url&quot;:&quot;&quot;,&quot;cover&quot;:&quot;http://mmbiz.qpic.cn/mmbiz_jpg/Z2f4yGjIu86oGv9lFn90JIEto9dzuPu00RlyWA8CbRYibD2zfFUo7aC4sYBViaOW3EwQTQwGCPRmjO04icZbybfLQ/0?wx_fmt=jpeg&quot;,&quot;author&quot;:&quot;吴波&quot;,&quot;copyright_stat&quot;:11,&quot;del_flag&quot;:1,&quot;item_show_type&quot;:0,&quot;audio_fileid&quot;:0,&quot;duration&quot;:0,&quot;play_url&quot;:&quot;&quot;,&quot;malicious_title_reason_id&quot;:0,&quot;malicious_content_type&quot;:0}],&quot;author&quot;:&quot;王士锋&quot;,&quot;copyright_stat&quot;:11,&quot;duration&quot;:0,&quot;del_flag&quot;:1,&quot;item_show_type&quot;:0,&quot;audio_fileid&quot;:0,&quot;play_url&quot;:&quot;&quot;,&quot;malicious_title_reason_id&quot;:0,&quot;malicious_content_type&quot;:0}},{&quot;comm_msg_info&quot;:{&quot;id&quot;:1000000021,&quot;type&quot;:49,&quot;datetime&quot;:1582379307,&quot;fakeid&quot;:&quot;3231769612&quot;,&quot;status&quot;:2,&quot;content&quot;:&quot;&quot;},&quot;app_msg_ext_info&quot;:{&quot;title&quot;:&quot;“防疫不停工，你我共学习”第一期：信号完整性基础理论讲解&quot;,&quot;digest&quot;:&quot;做好疫情防御的同时科研生产不能停，提升自己不能停，事业部推出“防疫不停工，你我共学习”在线学习活动。&quot;,&quot;content&quot;:&quot;&quot;,&quot;fileid&quot;:100000564,&quot;content_url&quot;:&quot;http://mp.weixin.qq.com/s?__biz=MzIzMTc2OTYxMg==&amp;amp;mid=2247484231&amp;amp;idx=1&amp;amp;sn=31e375471605a85144aaa05bd329014d&amp;amp;chksm=e89e5024dfe9d932f3bb7a8d4c3b0ace0556d595cade743aac84bd9c59b0ff787b8ae52e0df9&amp;amp;scene=27#wechat_redirect&quot;,&quot;source_url&quot;:&quot;&quot;,&quot;cover&quot;:&quot;http://mmbiz.qpic.cn/mmbiz_jpg/Z2f4yGjIu85phju9qwVUibP5kaLrGnJy6H5EFicnF7XiaDxu4gcqScRDCTGZz3csPxwyAWEAGH95ibMWnKrIIicC0sA/0?wx_fmt=jpeg&quot;,&quot;subtype&quot;:9,&quot;is_multi&quot;:1,&quot;multi_app_msg_item_list&quot;:[{&quot;title&quot;:&quot;“防疫不停工，你我共学习”第二期：信号完整性应用实践&quot;,&quot;digest&quot;:&quot;做好疫情防御的同时科研生产不能停，提升自己不能停，事业部推出“防疫不停工，你我共学习”在线学习活动。&quot;,&quot;content&quot;:&quot;&quot;,&quot;fileid&quot;:0,&quot;content_url&quot;:&quot;http://mp.weixin.qq.com/s?__biz=MzIzMTc2OTYxMg==&amp;amp;mid=2247484231&amp;amp;idx=2&amp;amp;sn=e64780c7270c3b21c063b46f2c53ca79&amp;amp;chksm=e89e5024dfe9d932aeb02ac1721ba6aaa783b71575a46a40f8270a36ec5f375a69ccc2d955b8&amp;amp;scene=27#wechat_redirect&quot;,&quot;source_url&quot;:&quot;&quot;,&quot;cover&quot;:&quot;http://mmbiz.qpic.cn/mmbiz_jpg/Z2f4yGjIu85phju9qwVUibP5kaLrGnJy6S9XLuqEq4OFeIacoicFfpreBic32sjibJVnMRNNibSRmYK9LVD24uLgvSg/0?wx_fmt=jpeg&quot;,&quot;author&quot;:&quot;二部一组&quot;,&quot;copyright_stat&quot;:11,&quot;del_flag&quot;:1,&quot;item_show_type&quot;:0,&quot;audio_fileid&quot;:0,&quot;duration&quot;:0,&quot;play_url&quot;:&quot;&quot;,&quot;malicious_title_reason_id&quot;:0,&quot;malicious_content_type&quot;:0},{&quot;title&quot;:&quot;“防疫不停工，你我共学习”第三期：高性能DSP专题技术交流&quot;,&quot;digest&quot;:&quot;做好疫情防御的同时科研生产不能停，提升自己不能停，事业部推出“防疫不停工，你我共学习”在线学习活动。&quot;,&quot;content&quot;:&quot;&quot;,&quot;fileid&quot;:0,&quot;content_url&quot;:&quot;http://mp.weixin.qq.com/s?__biz=MzIzMTc2OTYxMg==&amp;amp;mid=2247484231&amp;amp;idx=3&amp;amp;sn=ec0359c4cf38b535388abf6f11a06cb6&amp;amp;chksm=e89e5024dfe9d932c48523f42a81cd413605e21eb0afa8bcfefdcf9e6d5f5b26193cfedde6c2&amp;amp;scene=27#wechat_redirect&quot;,&quot;source_url&quot;:&quot;&quot;,&quot;cover&quot;:&quot;http://mmbiz.qpic.cn/mmbiz_jpg/Z2f4yGjIu85phju9qwVUibP5kaLrGnJy6TrQmf9zVVSgeovEPH8Dx9ECNwdNsc8tmWeXKlPXeq2fxeUIA6ZZv2g/0?wx_fmt=jpeg&quot;,&quot;author&quot;:&quot;二部一组&quot;,&quot;copyright_stat&quot;:11,&quot;del_flag&quot;:1,&quot;item_show_type&quot;:0,&quot;audio_fileid&quot;:0,&quot;duration&quot;:0,&quot;play_url&quot;:&quot;&quot;,&quot;malicious_title_reason_id&quot;:0,&quot;malicious_content_type&quot;:0}],&quot;author&quot;:&quot;二部一组&quot;,&quot;copyright_stat&quot;:11,&quot;duration&quot;:0,&quot;del_flag&quot;:1,&quot;item_show_type&quot;:0,&quot;audio_fileid&quot;:0,&quot;play_url&quot;:&quot;&quot;,&quot;malicious_title_reason_id&quot;:0,&quot;malicious_content_type&quot;:0}},{&quot;comm_msg_info&quot;:{&quot;id&quot;:1000000020,&quot;type&quot;:49,&quot;datetime&quot;:1582338209,&quot;fakeid&quot;:&quot;3231769612&quot;,&quot;status&quot;:2,&quot;content&quot;:&quot;&quot;},&quot;app_msg_ext_info&quot;:{&quot;title&quot;:&quot;“疫”呼百应&nbsp;|&nbsp;夯实战斗堡垒，打赢阻疫之战&quot;,&quot;digest&quot;:&quot;综合管理办在预防新冠肺炎的工作中，积极制定了实施策划，保证每日值班人员处理日常事务，日常操作流程落实到文件，保证不放过任何一个风险细节。&quot;,&quot;content&quot;:&quot;&quot;,&quot;fileid&quot;:100000531,&quot;content_url&quot;:&quot;http://mp.weixin.qq.com/s?__biz=MzIzMTc2OTYxMg==&amp;amp;mid=2247484194&amp;amp;idx=1&amp;amp;sn=c035b75181c7eef8e799d4b9a4702513&amp;amp;chksm=e89e5041dfe9d957b242652a593d457ac6fd5d8dc620751780a10660952dfe2c5b23fbc6eb8a&amp;amp;scene=27#wechat_redirect&quot;,&quot;source_url&quot;:&quot;&quot;,&quot;cover&quot;:&quot;http://mmbiz.qpic.cn/mmbiz_jpg/Z2f4yGjIu87IPYezXPZXWnP4CHQoLE2AkbG9PYibY4O3TuOXmtVJHDPVIfHvA3M4qSIo1zj8ciaBVFkuotvJowxg/0?wx_fmt=jpeg&quot;,&quot;subtype&quot;:9,&quot;is_multi&quot;:0,&quot;multi_app_msg_item_list&quot;:[],&quot;author&quot;:&quot;综合管理办&quot;,&quot;copyright_stat&quot;:11,&quot;duration&quot;:0,&quot;del_flag&quot;:1,&quot;item_show_type&quot;:0,&quot;audio_fileid&quot;:0,&quot;play_url&quot;:&quot;&quot;,&quot;malicious_title_reason_id&quot;:0,&quot;malicious_content_type&quot;:0}},{&quot;comm_msg_info&quot;:{&quot;id&quot;:1000000019,&quot;type&quot;:49,&quot;datetime&quot;:1582164339,&quot;fakeid&quot;:&quot;3231769612&quot;,&quot;status&quot;:2,&quot;content&quot;:&quot;&quot;},&quot;app_msg_ext_info&quot;:{&quot;title&quot;:&quot;打赢战“疫”&nbsp;|&nbsp;科学防疫，有序复工&quot;,&quot;digest&quot;:&quot;系统产品事业部在所党委和所疫情防控办公室的领导下，成立联防联控指挥中心，确保节后复工有序进行。&quot;,&quot;content&quot;:&quot;&quot;,&quot;fileid&quot;:100000529,&quot;content_url&quot;:&quot;http://mp.weixin.qq.com/s?__biz=MzIzMTc2OTYxMg==&amp;amp;mid=2247484187&amp;amp;idx=1&amp;amp;sn=4797c7d82e929ff52117a014368985c0&amp;amp;chksm=e89e5078dfe9d96e355c4e066708937a9f406fcd03fa0ca455907b61631652a87d45bd880c32&amp;amp;scene=27#wechat_redirect&quot;,&quot;source_url&quot;:&quot;&quot;,&quot;cover&quot;:&quot;http://mmbiz.qpic.cn/mmbiz_jpg/Z2f4yGjIu84gU0BTP0xVicdJHWYh0y77MCtu3WGPhNlBgqMYCCFpV9O4xGZBcYKxHNw3OWQCIjDX0m2JQ0uXO6A/0?wx_fmt=jpeg&quot;,&quot;subtype&quot;:9,&quot;is_multi&quot;:0,&quot;multi_app_msg_item_list&quot;:[],&quot;author&quot;:&quot;张红&quot;,&quot;copyright_stat&quot;:11,&quot;duration&quot;:0,&quot;del_flag&quot;:1,&quot;item_show_type&quot;:0,&quot;audio_fileid&quot;:0,&quot;play_url&quot;:&quot;&quot;,&quot;malicious_title_reason_id&quot;:0,&quot;malicious_content_type&quot;:0}},{&quot;comm_msg_info&quot;:{&quot;id&quot;:1000000018,&quot;type&quot;:49,&quot;datetime&quot;:1581259369,&quot;fakeid&quot;:&quot;3231769612&quot;,&quot;status&quot;:2,&quot;content&quot;:&quot;&quot;},&quot;app_msg_ext_info&quot;:{&quot;title&quot;:&quot;系统产品事业部在新冠肺炎疫情期间复工上岗总要求&quot;,&quot;digest&quot;:&quot;疫情防控期间复工上岗总要求&quot;,&quot;content&quot;:&quot;&quot;,&quot;fileid&quot;:100000516,&quot;content_url&quot;:&quot;http://mp.weixin.qq.com/s?__biz=MzIzMTc2OTYxMg==&amp;amp;mid=2247484168&amp;amp;idx=1&amp;amp;sn=31e5699567fe45e878d117343fc42281&amp;amp;chksm=e89e506bdfe9d97d6631c7a418e69d21edcb7c6e6f49c3f3e17abc03ce5dea5531bf67fe8bca&amp;amp;scene=27#wechat_redirect&quot;,&quot;source_url&quot;:&quot;&quot;,&quot;cover&quot;:&quot;http://mmbiz.qpic.cn/mmbiz_jpg/Z2f4yGjIu85v4Y7cHFzKaHUVeJ4noKJFAgkPP3rRdSLAbMcRMUhdhpUpLfe1icyB39441ssvm603ZhNAneVFTXA/0?wx_fmt=jpeg&quot;,&quot;subtype&quot;:9,&quot;is_multi&quot;:0,&quot;multi_app_msg_item_list&quot;:[],&quot;author&quot;:&quot;系统产品事业部&quot;,&quot;copyright_stat&quot;:11,&quot;duration&quot;:0,&quot;del_flag&quot;:1,&quot;item_show_type&quot'''
##
####s = html.unescape(s).replace('\xa0', ' ')
####res = re.findall(r'"title":"(.*?)".*?"content_url":"(.*?)"', s)

with open('历史文章.html', encoding='u8') as f:
    s = f.read()

urls = re.findall(r'<h4 class="weui_media_title" hrefs="(.*?)">', s)

for url in urls:
    try:
        SavePassage(url)
    except Exception as e:
        print(e)
        print(url)


